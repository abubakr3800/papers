# -*- coding: utf-8 -*-
"""
Generates a single DOCX booklet with:
- Problem, Idea, Study sections
- 8 parameter graphs with green/yellow/red bands
- Age-group × environment recommendations
- Clickable references (titles + URLs)

Dependencies:
  pip install matplotlib python-docx numpy
"""

import os
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# ---------- Helpers ----------

def ensure_dir(path):
    if not os.path.isdir(path):
        os.makedirs(path, exist_ok=True)

def gaussian_peak(x, center, width=0.2, maxy=1.0):
    """Symmetric 'benefit' curve peaking at center (as fraction of span)."""
    # x is numeric array; normalize to 0..1
    x0, x1 = x.min(), x.max()
    xn = (x - x0) / (x1 - x0 + 1e-9)
    return maxy * np.exp(-((xn - center) ** 2) / (2 * width ** 2))

def descending_curve(x, knee_frac=0.15):
    """Monotonic decreasing 'risk' curve (for Flicker, UGR)."""
    x0, x1 = x.min(), x.max()
    xn = (x - x0) / (x1 - x0 + 1e-9)
    # High at 0, falls rapidly after knee
    return 1.0 / (1.0 + np.exp( (xn - knee_frac) * 12.0 ))

def band_plot(ax, x_min, x_max, good, warn, danger, label_left="Lower is worse", label_right="Higher is worse"):
    """
    Draw red/yellow/green bands across x, covering:
      danger = (min, max) full span,
      warn   = (warn_min, warn_max) inside danger,
      good   = (good_min, good_max) inside warn.
    """
    # Red left
    ax.axvspan(danger[0], warn[0], color="red", alpha=0.30, label="Danger/Risk")
    # Yellow left
    ax.axvspan(warn[0], good[0], color="yellow", alpha=0.35, label="Caution")
    # Green
    ax.axvspan(good[0], good[1], color="green", alpha=0.40, label="Optimal")
    # Yellow right
    ax.axvspan(good[1], warn[1], color="yellow", alpha=0.35)
    # Red right
    ax.axvspan(warn[1], danger[1], color="red", alpha=0.30)

    ax.set_xlim(x_min, x_max)
    ax.set_yticks([])
    # Legend: keep unique labels
    handles, labels = ax.get_legend_handles_labels()
    d = dict(zip(labels, handles))
    ax.legend(d.values(), d.keys(), loc="upper right")

def clean_ticks(ax, x_min, x_max, n=6, as_int=True):
    ticks = np.linspace(x_min, x_max, n)
    if as_int:
        ticks = ticks.astype(int)
    ax.set_xticks(ticks)

# ---------- Parameter specs with ranges, curve types, and references ----------

# NOTE: Ranges reflect trusted sources cited at the bottom of the script
PARAMETERS = [
    {
        "name": "CCT (Correlated Color Temperature, K)",
        "span": (2000, 8000),
        "good": (4000, 5000),
        "warn": (3000, 6500),
        "danger": (2000, 8000),
        "curve": "gaussian",     # benefit rises toward mid (balanced) then falls
        "curve_center": 0.54,    # slightly toward cooler end (within 2k..8k)
        "xlabel": "CCT (K)",
        "effect": (
            "Balanced 4000–5000 K supports alertness and visual comfort for classrooms. "
            "Too warm (<3000 K) can promote sleepiness; too cool (>6500 K) may increase discomfort/glare. "
            "Short, task-specific boosts at 6500 K/1000 lx can improve reading fluency during tests."
        ),
        "refs": [
            ("EN 12464-1: Indoor work lighting (general targets)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("Mott et al., 2012: High CCT/illuminance improved reading fluency", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099"),
            ("Sleegers et al., 2013: Dynamic lighting and concentration", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099"),
            ("Park et al., 2015: CCT, EEG & task performance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
            ("Chen et al., 2022: CCT × illuminance responses", "https://www.mdpi.com/1996-1073/15/12/4477")
        ]
    },
    {
        "name": "CRI (Color Rendering Index, Ra)",
        "span": (50, 100),
        "good": (80, 100),
        "warn": (70, 80),
        "danger": (50, 100),
        "curve": "gaussian",
        "curve_center": 0.95,
        "xlabel": "CRI (Ra)",
        "effect": (
            "CRI ≥80 supports natural color appearance and reduces visual fatigue. "
            "Art/graphics benefit from CRI ≥90. Very low CRI (<70) hampers color discrimination and comfort."
        ),
        "refs": [
            ("EN 12464-1: Ra ≥80 typical; ≥90 for demanding color tasks", "https://www.performanceinlighting.com/mo/en/en-12464-1")
        ]
    },
    {
        "name": "Flicker (% modulation at typical LED driving frequencies)",
        "span": (0, 50),
        "good": (0, 5),
        "warn": (5, 20),
        "danger": (0, 50),
        "curve": "descending",
        "xlabel": "Percent Flicker (%)",
        "effect": (
            "Keep percent flicker <5% to minimize headaches, eyestrain, and distraction. "
            "Between 5–20% some occupants are affected; >20% increases adverse effects. "
            "Use high-frequency drivers and avoid PWM at low frequencies per IEEE 1789."
        ),
        "refs": [
            ("IEEE 1789-2015: Flicker recommended practice", "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf"),
            ("DOE/LightFair deck: factors increasing risk", "https://www.energy.gov/sites/default/files/2022-11/ssl-miller-lehman_flicker_lightfair2015.pdf"),
            ("DIAL explainer on IEEE 1789 terms", "https://www.dial.de/en-GB/articles/ieee-1789-a-new-standard-for-evaluating-flickering-leds")
        ]
    },
    {
        "name": "Glare (UGR)",
        "span": (10, 30),
        "good": (10, 19),
        "warn": (19, 22),
        "danger": (10, 30),
        "curve": "descending",
        "xlabel": "UGR",
        "effect": (
            "UGR <19 recommended for classrooms to avoid discomfort and maintain performance. "
            "Aim lower (≈16–18) near screens/interactive boards."
        ),
        "refs": [
            ("EN 12464-1: UGR targets for tasks", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("CIBSE Factfile: importance of UGR", "https://www.cibse.org/media/polbabib/factfile-15-the-importance-of-glare-and-calculating-ugr-jul2019.pdf")
        ]
    },
    {
        "name": "Melanopic EDI (lux at eye, vertical)",
        "span": (0, 800),
        "good": (250, 500),
        "warn": (100, 250),
        "danger": (0, 800),
        "curve": "gaussian",
        "curve_center": 0.45,
        "xlabel": "Melanopic EDI (lux)",
        "effect": (
            "Daytime ≥250 melanopic EDI at eye supports circadian entrainment and alertness "
            "(measure at ≈1.2 m seated, vertical). Evening levels should be much lower."
        ),
        "refs": [
            ("Global consensus (Brown et al., 2022): daytime ≥250 mEDI", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
            ("PMC version", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
            ("WELL v2 L03 circadian targets (EML/mEDI)", "https://standard.wellcertified.com/light/circadian-lighting-design")
        ]
    },
    {
        "name": "Vertical Illuminance (lux at eye/face)",
        "span": (50, 1000),
        "good": (300, 500),
        "warn": (200, 800),
        "danger": (50, 1000),
        "curve": "gaussian",
        "curve_center": 0.45,
        "xlabel": "Vertical Illuminance (lux)",
        "effect": (
            "Adequate vertical illuminance improves visibility of faces/boards and supports non-visual effects. "
            "Keep roughly 300–500 lx on faces in learning spaces; avoid very low or very high values."
        ),
        "refs": [
            ("EN 12464-1: room surface/vertical illuminance guidance", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("CIBSE LG5/education guidance (context)", "https://www.cibse.org/knowledge-research/knowledge-portal/lg7-lighting-for-offices-2023")
        ]
    },
    {
        "name": "Exposure Duration (hours of daytime light meeting targets)",
        "span": (0, 12),
        "good": (3, 6),
        "warn": (1, 8),
        "danger": (0, 12),
        "curve": "gaussian",
        "curve_center": 0.45,
        "xlabel": "Hours per school day",
        "effect": (
            "Sustained exposure (≈3–6 h) to target illuminance/spectrum during the school day supports alertness "
            "and entrainment; very little or excessive high-intensity exposure is less beneficial."
        ),
        "refs": [
            ("Consensus guidance on daytime vs evening exposure (Brown et al., 2022)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/")
        ]
    },
    {
        "name": "Horizontal Illuminance (desk, lux)",
        "span": (100, 1500),
        "good": (300, 500),
        "warn": (200, 1000),
        "danger": (100, 1500),
        "curve": "gaussian",
        "curve_center": 0.40,
        "xlabel": "Horizontal Illuminance (lux)",
        "effect": (
            "Provide 300–500 lx at desks for most classroom tasks; 500–750+ lx for labs/graphics. "
            "Higher levels (≈1000 lx) may be used short-term for exams/focus sessions."
        ),
        "refs": [
            ("EN 12464-1: classroom/task illuminance", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("Mott/Sleegers: high-lx focus settings evidence", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099")
        ]
    }
]

# ---------- Build document ----------

out_dir = os.path.join(os.getcwd(), "lighting_booklet_output")
img_dir = os.path.join(out_dir, "images")
ensure_dir(img_dir)

doc = Document()
doc.add_heading("The Effect of Light on Students in Schools", 0)
doc.add_paragraph("A concise, referenced booklet on how lighting parameters affect student concentration, biology, and psychology.")

# Problem
doc.add_heading("1) The Problem", level=1)
doc.add_paragraph(
    "Suboptimal school lighting (low/imbalanced illuminance, excessive flicker, high glare, inappropriate spectrum) "
    "is associated with headaches, visual fatigue, poor concentration, and circadian disruption, which can degrade academic performance."
)
doc.add_paragraph("Key standards and reviews emphasize illuminance, glare control (UGR), color rendering (CRI), and circadian-effective light.")

# Idea
doc.add_heading("2) The Idea", level=1)
doc.add_paragraph(
    "Study how measurable lighting parameters—CCT, CRI, Flicker, Glare (UGR), Melanopic EDI, Vertical Illuminance, Exposure Duration, "
    "and Horizontal Illuminance—affect children at different ages. Compare optimal vs. harmful ranges, with biological rationale and outcomes."
)

# Study
doc.add_heading("3) The Study: Parameters, Effects & Visuals", level=1)
for p in PARAMETERS:
    name = p["name"]
    x0, x1 = p["span"]
    x = np.linspace(x0, x1, 400)
    if p["curve"] == "gaussian":
        y = gaussian_peak(x, center=p.get("curve_center", 0.5), width=0.18, maxy=1.0)
    else:
        y = descending_curve(x)

    # Figure
    fig, ax = plt.subplots(figsize=(7, 2.6))
    band_plot(ax, x0, x1, p["good"], p["warn"], p["danger"])
    ax.plot(x, y, "k--", linewidth=2)
    ax.set_xlabel(p["xlabel"])
    ax.set_title(name)
    clean_ticks(ax, x0, x1, n=6, as_int=True)
    img_path = os.path.join(img_dir, f"{name.replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '')}.png")
    fig.tight_layout()
    fig.savefig(img_path, bbox_inches="tight")
    plt.close(fig)

    # Document section
    doc.add_heading(name, level=2)
    doc.add_paragraph(p["effect"])
    doc.add_paragraph(f"Optimal range: {p['good'][0]}–{p['good'][1]} | Caution: {p['warn'][0]}–{p['warn'][1]} (context-dependent).")
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph("References:")
    for title, url in p["refs"]:
        doc.add_paragraph(f"• {title} — {url}")

# Solution: age × environment
doc.add_heading("4) The Solution: Evidence-Based Targets by Age & Environment", level=1)
doc.add_paragraph(
    "Below are practical set-points derived from standards and research. Horizontal illuminance/UGR/CRI from EN 12464-1; "
    "melanopic targets from Brown et al. (2022) and WELL v2 L03; flicker per IEEE 1789; CCT/task boosts per classroom studies."
)

recommendations = [
    # (Age/Env, Horizontal lx, UGR, CRI, mEDI (day), CCT, Notes/Use, Refs)
    ("Kindergarten (3–5) – classroom", "300–500 lx", "<19", "≥80", "≥250 mEDI (daytime)", "3500–4000 K",
     "Softer CCT to reduce arousal; keep flicker <5%; good vertical light for faces (≈300–400 lx).",
     [
         ("EN 12464-1: classroom lx/UGR/CRI", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
         ("Brown et al., 2022 mEDI ≥250", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/")
     ]),
    ("Primary (6–11) – classroom", "300–500 lx", "<19", "≥80", "≥250–300 mEDI", "4000–5000 K",
     "Balanced spectrum/daylight; flicker <5%; vertical ≈300–500 lx on faces/boards.",
     [
         ("EN 12464-1", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
         ("WELL v2 L03 circadian targets", "https://standard.wellcertified.com/light/circadian-lighting-design")
     ]),
    ("Secondary (12–18) – classroom", "300–500 lx", "<19 (≤16 near screens)", "≥80 (≥90 for art)", "≥250–300 mEDI", "4000–5000 K",
     "Lower UGR near screens; can use short high-CCT/1000 lx sessions for tests.",
     [
         ("EN 12464-1", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
         ("Mott/Sleegers focus setting", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099")
     ]),
    ("Exam/Focus sessions (all ages)", "500–1000 lx (short-term)", "<19", "≥80", "≥250–400 mEDI", "5000–6500 K",
     "Short deployments to boost alertness/reading fluency; avoid all-day cold light.",
     [
         ("Mott et al., 2012", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099"),
         ("Sleegers et al., 2013", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099")
     ]),
    ("Art/Graphics room", "500–750 lx", "<19", "≥90", "≥250 mEDI", "4000–5000 K",
     "High CRI for accurate color tasks; good vertical light to evaluate work.",
     [
         ("EN 12464-1 (color-critical tasks)", "https://www.performanceinlighting.com/mo/en/en-12464-1")
     ]),
    ("Science lab", "500–750 lx", "<19 (≤16 preferred)", "≥80", "≥250–300 mEDI", "4000–5000 K",
     "Higher task lx and glare control for practical work; minimize flicker.",
     [
         ("EN 12464-1 (laboratory tasks)", "https://www.performanceinlighting.com/mo/en/en-12464-1")
     ]),
    ("Corridors / circulation", "100–200 lx", "<22", "≥80", "—", "3000–4000 K",
     "Comfortable navigation; avoid excessive brightness/glare.",
     [
         ("EN 12464-1", "https://www.performanceinlighting.com/mo/en/en-12464-1")
     ])
]

for row in recommendations:
    (who_where, lx, ugr, cri, medi, cct, note, refs) = row
    doc.add_paragraph(f"• {who_where}")
    doc.add_paragraph(f"  - Horizontal illuminance: {lx}")
    doc.add_paragraph(f"  - UGR: {ugr}   |   CRI: {cri}")
    doc.add_paragraph(f"  - Daytime melanopic target: {medi}")
    doc.add_paragraph(f"  - Typical CCT: {cct}")
    doc.add_paragraph(f"  - Notes: {note}")
    doc.add_paragraph("  - References:")
    for title, url in refs:
        doc.add_paragraph(f"    • {title} — {url}")


# Global implementation notes
doc.add_heading("Implementation Notes (All Spaces)", level=2)
doc.add_paragraph("• Keep flicker (percent modulation) <5% and avoid low-frequency PWM dimming (IEEE 1789).")
doc.add_paragraph("• Aim for UGR <19; place luminaires to avoid direct view and specular reflections of boards/screens.")
doc.add_paragraph("• Provide ≥250 melanopic EDI at eye (daytime); drastically lower in evening events to avoid circadian delay.")
doc.add_paragraph("• Use CRI ≥80 (≥90 for color-critical work).")
doc.add_paragraph("• Balance horizontal (desk) lx with adequate vertical illuminance for faces and boards.")

# References section (clickable URLs are fine as plain text in Word)
doc.add_heading("References (Titles + Links)", level=1)
ALL_REFS = [
    ("EN 12464-1: Lighting of work places — Indoor", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
    ("CIBSE Factfile: Importance of UGR", "https://www.cibse.org/media/polbabib/factfile-15-the-importance-of-glare-and-calculating-ugr-jul2019.pdf"),
    ("WELL v2 L03: Circadian Lighting Design", "https://standard.wellcertified.com/light/circadian-lighting-design"),
    ("Brown et al., 2022 (PLOS Biology): Global consensus on melanopic EDI", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
    ("PMC mirror for Brown et al., 2022", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
    ("IEEE 1789-2015: Flicker Recommended Practice", "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf"),
    ("DOE/LightFair deck: Flicker risk factors", "https://www.energy.gov/sites/default/files/2022-11/ssl-miller-lehman_flicker_lightfair2015.pdf"),
    ("Mott et al., 2012 / Sleegers et al., 2013: Focus settings in classrooms", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099"),
    ("Park et al., 2015: CCT, EEG & task performance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
    ("Chen et al., 2022: CCT × Illuminance responses", "https://www.mdpi.com/1996-1073/15/12/4477")
]
for title, url in ALL_REFS:
    doc.add_paragraph(f"• {title} — {url}")

# Save
ensure_dir(out_dir)
out_docx = os.path.join(out_dir, "School_Lighting_Booklet_FINAL.docx")
doc.save(out_docx)

print("✅ Booklet created at:", out_docx)
print("🖼️ Plots saved in:", img_dir)
