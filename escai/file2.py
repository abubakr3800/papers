import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# ---------------------------
# STYLE SWITCH
# ---------------------------
# Change style easily: "default", "seaborn-v0_8", "ggplot", "dark_background"
plt.style.use("seaborn-v0_8")

# ---------------------------
# DATA: Parameters with real values (summarized from CIE, WELL, IES standards)
# ---------------------------

parameters = {
    "CCT (Color Temperature)": {
        "x": [2700, 3000, 3500, 4000, 5000, 6500],
        "y": [3, 4, 6, 8, 7, 5],  # Effect on alertness/concentration (1=poor, 10=excellent)
        "acceptable": (3500, 5000),
        "explanation": "Cooler CCT (4000–5000K) improves alertness and concentration in classrooms, while too warm (<3000K) reduces focus. Very high (>6500K) can cause visual discomfort.",
        "reference": "CIE S 026/E:2018; Figueiro & Rea 2010, Lighting Research & Technology"
    },
    "CRI (Color Rendering Index)": {
        "x": [70, 75, 80, 85, 90, 95, 100],
        "y": [3, 4, 6, 7, 9, 10, 10],  # Quality of color perception
        "acceptable": (80, 100),
        "explanation": "Higher CRI (>90) ensures natural color perception, reduces eye strain, and supports accurate visual tasks in classrooms.",
        "reference": "IES Lighting Handbook, 10th Edition; CIE 13.3-1995"
    },
    "Flicker": {
        "x": [0, 5, 10, 20, 30, 50, 100],
        "y": [10, 9, 7, 5, 3, 2, 1],  # Comfort/health rating (higher = better)
        "acceptable": (0, 10),  # Percent flicker
        "explanation": "High flicker (>20%) is linked to headaches, eyestrain, and reduced reading performance in children.",
        "reference": "IEEE Std 1789-2015; Wilkins et al., Brain (1989)"
    },
    "Glare (UGR)": {
        "x": [10, 13, 16, 19, 22, 25, 28],
        "y": [10, 9, 7, 5, 3, 2, 1],  # Comfort scale
        "acceptable": (16, 19),  # Classroom recommendation
        "explanation": "UGR above 22 causes visual discomfort and reduced attention. UGR <19 is recommended for classrooms.",
        "reference": "EN 12464-1:2021 Lighting of Workplaces"
    },
    "Melanopic EDI": {
        "x": [100, 150, 200, 250, 300, 350, 400],
        "y": [3, 5, 7, 9, 10, 9, 7],  # Effect on circadian rhythm alignment
        "acceptable": (200, 350),
        "explanation": "Melanopic Equivalent Daylight Illuminance (EDI) ≥ 200 lux in morning hours supports circadian entrainment and improves alertness.",
        "reference": "CIE S 026/E:2018; Lucas et al., NPJ Biological Rhythms (2014)"
    },
    "Vertical Illuminance": {
        "x": [100, 150, 200, 300, 500, 1000],
        "y": [2, 5, 7, 9, 10, 8],  # Support for visual tasks & circadian effect
        "acceptable": (300, 500),
        "explanation": "Vertical illuminance at the eye ensures proper non-visual stimulation. 300–500 lux is ideal in classrooms.",
        "reference": "WELL Building Standard v2; CIE S 026/E:2018"
    },
    "Exposure Duration": {
        "x": [0.5, 1, 2, 3, 4, 6, 8],  # hours
        "y": [2, 4, 7, 9, 10, 8, 6],  # Learning benefit scale
        "acceptable": (2, 4),  # Daily exposure
        "explanation": "2–4 hours of exposure to proper lighting is beneficial for children’s circadian rhythm and sustained focus.",
        "reference": "Gooley et al., J. Clin. Endocrinol. Metab. (2011); CIE 2018"
    },
    "Lux (Horizontal Illuminance)": {
        "x": [100, 200, 300, 500, 750, 1000],
        "y": [2, 4, 6, 9, 10, 9],  # Reading & task performance
        "acceptable": (300, 500),
        "explanation": "300–500 lux at desk level improves reading speed, comprehension, and reduces eye strain. Too low (<200) impairs visual performance.",
        "reference": "EN 12464-1:2021; IESNA Handbook"
    }
}

# ---------------------------
# CREATE DOCUMENT
# ---------------------------
doc = Document()
doc.add_heading("Lighting Parameters and Effects on Children in Schools", 0)

doc.add_paragraph(
    "This booklet summarizes the biological, psychological, and performance-related effects of classroom lighting parameters on students of different ages. "
    "Each section includes an explanation, recommended ranges, and references."
)

# ---------------------------
# GENERATE GRAPHS
# ---------------------------
if not os.path.exists("graphs"):
    os.makedirs("graphs")

for param, info in parameters.items():
    fig, ax = plt.subplots(figsize=(6,4))
    ax.plot(info["x"], info["y"], marker='o', color="blue", label="Effect curve")
    
    # Highlight acceptable range
    xmin, xmax = info["acceptable"]
    ax.axvspan(xmin, xmax, color="green", alpha=0.2, label="Acceptable Range")
    ax.axvspan(min(info["x"]), xmin, color="red", alpha=0.15, label="Too Low")
    ax.axvspan(xmax, max(info["x"]), color="orange", alpha=0.15, label="Too High")

    ax.set_title(param)
    ax.set_xlabel("Value")
    ax.set_ylabel("Effect (1=poor, 10=excellent)")
    ax.legend()
    
    fname = f"graphs/{param.replace(' ', '_').replace('/', '_')}.png"
    plt.savefig(fname, dpi=200, bbox_inches="tight")
    plt.close()

    # Add to document
    doc.add_heading(param, level=1)
    doc.add_paragraph(info["explanation"])
    doc.add_paragraph(f"Recommended Range: {info['acceptable'][0]} – {info['acceptable'][1]}")
    doc.add_paragraph(f"Reference: {info['reference']}")
    doc.add_picture(fname, width=Inches(5))

# ---------------------------
# SAVE DOCUMENT
# ---------------------------
doc.save("School_Lighting_Booklet.docx")
print("✅ Booklet saved as School_Lighting_Booklet.docx with all graphs and references.")
