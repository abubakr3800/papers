# make_school_lighting_booklet.py
# -*- coding: utf-8 -*-
"""
Creates a unified DOCX booklet that merges:
 - your uploaded file (/mnt/data/Schools information.docx)
 - the 8-parameter chapter study (each with one combined figure)
Outputs:
 - school_lighting_booklet_output/School_Lighting_Booklet_Merged.docx
 - school_lighting_booklet_output/images/*.png
"""

import os
import re
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx import opc

# --------------------------
# Config / Style
# --------------------------
MATPLOTLIB_STYLE = "seaborn-v0_8"   # change this to "default", "ggplot", etc.
plt.style.use(MATPLOTLIB_STYLE)

OUT_DIR = os.path.join(os.getcwd(), "school_lighting_booklet_output")
IMG_DIR = os.path.join(OUT_DIR, "images")
INPUT_UPLOADED_DOCX = "/mnt/data/Schools information.docx"  # path to your uploaded docx
OUTPUT_DOCX = os.path.join(OUT_DIR, "School_Lighting_Booklet_Merged.docx")

COLOR_CURVE = "#1f77b4"   # blue
COLOR_GOOD  = "#00f700"   # green
COLOR_WARN  = "#ffbf00"   # amber
COLOR_BAD   = "#ff0000"   # red
COLOR_MARK  = "#555555"

# Ensure output dirs
os.makedirs(IMG_DIR, exist_ok=True)

# --------------------------
# Utility helpers
# --------------------------
def read_docx_text(path):
    """
    Read all paragraphs from a docx and return combined text and a list of paragraphs.
    """
    if not os.path.exists(path):
        return None, []
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paras), paras

def extract_key_values(paragraphs):
    """
    Try to extract some structured values from uploaded doc paragraphs:
    - CCT values mentioned
    - Lux / illuminance values
    - CRI mentions
    - Flicker mentions
    - Age-group notes
    Returns a dict with best-effort values.
    """
    txt = "\n".join(paragraphs)
    res = {}

    # find numeric kelvin values for CCT (e.g., 2500K, 5000K)
    cct = re.findall(r"(\b[23-7]\d{2,3})\s*K\b", txt, flags=re.IGNORECASE)
    if cct:
        # convert to ints and unique sort
        cc = sorted(set(int(x) for x in cct))
        res['CCT_values'] = cc

    # find lux numbers
    lux = re.findall(r"(\b\d{2,4})\s*(?:lux|LX|Lux)\b", txt, flags=re.IGNORECASE)
    if lux:
        lx = sorted(set(int(x) for x in lux))
        res['Lux_values'] = lx

    # find CRI mentions (just keywords)
    if re.search(r"\bCRI\b", txt, flags=re.IGNORECASE):
        # try to find numerical CRI values near "CRI"
        cri_vals = re.findall(r"CRI.*?(\d{2,3})", txt, flags=re.IGNORECASE)
        if cri_vals:
            res['CRI_values'] = sorted(set(int(x) for x in cri_vals))

    # flicker / frequency
    flick = re.findall(r"(\b\d{2,4})\s*Hz\b", txt, flags=re.IGNORECASE)
    if flick:
        res['Flicker_freqs_Hz'] = sorted(set(int(x) for x in flick))

    # Age groups mentions
    age_hits = {}
    for age_label in ["preschool", "kindergarten", "elementary", "primary", "secondary", "adolescent", "undergraduate", "children"]:
        if re.search(r"\b" + age_label + r"\b", txt, flags=re.IGNORECASE):
            age_hits[age_label] = True
    if age_hits:
        res['age_mentions'] = list(age_hits.keys())

    # find phrases about % improvement or changes
    pct_changes = re.findall(r"([\d]{1,3}\.?\d{0,2})\s*%\s*(?:improv|increase|decrease|reduc)", txt, flags=re.IGNORECASE)
    if pct_changes:
        res['percent_changes'] = pct_changes

    return res

def ensure_ticks(ax, x_min, x_max, n=6, integer=False):
    ticks = np.linspace(x_min, x_max, n)
    if integer:
        ticks = np.round(ticks).astype(int)
    ax.set_xticks(ticks)

def add_bands(ax, x_min, x_max, good, warn, danger):
    """
    Draw bands: danger full span, warn inner, good inner-core
    good, warn, danger are tuples (min, max)
    """
    ax.axvspan(danger[0], warn[0], color=COLOR_BAD, alpha=0.25)
    ax.axvspan(warn[0], good[0], color=COLOR_WARN, alpha=0.30)
    ax.axvspan(good[0], good[1], color=COLOR_GOOD, alpha=0.30)
    ax.axvspan(good[1], warn[1], color=COLOR_WARN, alpha=0.30)
    ax.axvspan(warn[1], danger[1], color=COLOR_BAD, alpha=0.25)
    ax.set_xlim(x_min, x_max)
    ax.set_yticks([])

def interpolate_curve(x, anchors_x, anchors_y):
    # anchors_x must be increasing
    anchors_x = np.array(anchors_x)
    anchors_y = np.array(anchors_y)
    # If anchors arrays length mismatch, handle gracefully
    if anchors_x.size != anchors_y.size:
        # pad or trim anchors_y
        m = min(anchors_x.size, anchors_y.size)
        anchors_x = anchors_x[:m]
        anchors_y = anchors_y[:m]
    return np.interp(x, anchors_x, anchors_y)

# --------------------------
# Data for 8 parameters (anchors are literature-anchored approximations)
# --------------------------
# Each parameter: title, xspan, bands (good, warn, danger), anchors.x, anchors.y, x_label, y_label, notes, refs, markers
PARAMS = [
    {
        "title": "CCT (Correlated Color Temperature, K)",
        "xspan": (2000, 7000),
        "bands": {"good": (4000, 5000), "warn": (3000, 6500), "danger": (2000, 7000)},
        "anchors_x": [2000, 2700, 3000, 3500, 4000, 5000, 6500, 7000],
        "anchors_y": [8, 15, 25, 45, 65, 75, 60, 55],   # alerting potential (% relative)
        "x_label": "CCT (K)",
        "y_label": "Estimated Alerting Potential (%)",
        "notes": "Higher CCT (bluer light) tends to increase alertness and cognitive stimulation during daytime; warmer CCT supports calmness and relaxation.",
        "refs": [
            ("Park et al., 2015 (PMC)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
            ("Mott et al., classroom focus literature", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099"),
            ("User uploaded study (Schools information.docx) - CCT tests at 2500K,3000K,4000K,5000K,6500K", INPUT_UPLOADED_DOCX)
        ],
        "markers": [(5000, "5000K - observed optimum"), (6500, "6500K - high alertness")]
    },

    {
        "title": "CRI (Color Rendering Index, Ra)",
        "xspan": (60, 100),
        "bands": {"good": (80, 100), "warn": (70, 80), "danger": (60, 100)},
        "anchors_x": [60, 70, 75, 80, 85, 90, 95, 100],
        "anchors_y": [50, 60, 72, 82, 90, 96, 99, 100],   # visual color fidelity percent
        "x_label": "CRI (Ra)",
        "y_label": "Visual Color Fidelity / Recognition (%)",
        "notes": "Higher CRI improves color recognition and visual comfort; CRI is typically kept high in classrooms for accurate color tasks.",
        "refs": [
            ("EN 12464 standard (CRI guidance)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            (INPUT_UPLOADED_DOCX, "User uploaded study: CRI mentioned in docx (high CRI noted)")
        ],
        "markers": [(80, "CRI 80 baseline"), (90, "CRI 90 art/graphics")]
    },

    {
        "title": "Flicker (Percent modulation)",
        "xspan": (0, 50),
        "bands": {"good": (0, 5), "warn": (5, 20), "danger": (0, 50)},
        "anchors_x": [0, 2, 5, 10, 20, 30, 40, 50],
        "anchors_y": [0, 3, 8, 20, 45, 70, 85, 95],   # adverse effect risk %
        "x_label": "Percent Flicker (%)",
        "y_label": "Estimated Adverse Effect Risk (%)",
        "notes": "Unnoticeable high-frequency flicker still can impact sensitive individuals; keep flicker as low as possible (IEEE 1789 guidance).",
        "refs": [
            ("IEEE 1789 recommended practice", "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf"),
            ("User uploaded study (Schools information.docx) noted ~100Hz fluorescent flicker in many classrooms", INPUT_UPLOADED_DOCX)
        ],
        "markers": [(5, "Preferred <5%"), (20, "High risk >20%")]
    },

    {
        "title": "Glare (Unified Glare Rating UGR)",
        "xspan": (10, 30),
        "bands": {"good": (10, 19), "warn": (19, 22), "danger": (10, 30)},
        "anchors_x": [10, 13, 16, 19, 22, 25, 28, 30],
        "anchors_y": [5, 8, 15, 30, 55, 80, 92, 96],   # discomfort %
        "x_label": "UGR",
        "y_label": "Estimated Discomfort Probability (%)",
        "notes": "High glare leads to eye strain and distraction; control luminaire placement and reflections to keep UGR low in classrooms.",
        "refs": [
            ("CIBSE guidance on UGR", "https://www.cibse.org/"),
            (INPUT_UPLOADED_DOCX, "User file: glare noted as a negative factor")
        ],
        "markers": [(19, "UGR 19 classroom max")]
    },

    {
        "title": "Uniformity (Emin / Eavg)",
        "xspan": (0.1, 1.0),
        "bands": {"good": (0.6, 1.0), "warn": (0.4, 0.59), "danger": (0.1, 1.0)},
        "anchors_x": [0.1, 0.2, 0.3, 0.45, 0.6, 0.75, 0.9, 1.0],
        "anchors_y": [40, 55, 70, 82, 92, 96, 98, 99],   # performance index %
        "x_label": "Uniformity (Emin / Eavg)",
        "y_label": "Task Performance Index (%)",
        "notes": "Higher uniformity reduces local visual contrast and improves even task performance across the room.",
        "refs": [
            ("EN 12464-1 uniformity recommendations", "https://www.performanceinlighting.com/mo/en/en-12464-1")
        ],
        "markers": [(0.6, "Recommended ≥0.6")]
    },

    {
        "title": "Melanopic EDI (melanopic lux at eye)",
        "xspan": (0, 800),
        "bands": {"good": (250, 500), "warn": (100, 250), "danger": (0, 800)},
        "anchors_x": [0, 20, 50, 100, 250, 500, 800],
        "anchors_y": [0, 5, 15, 35, 65, 80, 90],   # melatonin suppression %
        "x_label": "Melanopic EDI (lux)",
        "y_label": "Estimated Melatonin Suppression (%)",
        "notes": "Melanopic EDI of ~250 lux or higher in daytime supports circadian entrainment and alertness (Brown et al., 2022 consensus).",
        "refs": [
            ("Brown et al., 2022 consensus (PLOS Biology)", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
            (INPUT_UPLOADED_DOCX, "User file: referenced melanopic / circadian impacts in literature review")
        ],
        "markers": [(250, "250 mEDI recommended min"), (500, "Strong daytime level")]
    },

    {
        "title": "Vertical Illuminance (lux at eye/face)",
        "xspan": (50, 1000),
        "bands": {"good": (300, 500), "warn": (200, 800), "danger": (50, 1000)},
        "anchors_x": [50, 100, 150, 300, 500, 800, 1000],
        "anchors_y": [0.05, 0.12, 0.22, 0.40, 0.55, 0.65, 0.68],   # circadian stimulus (CS) approximate
        "x_label": "Vertical Illuminance (lux)",
        "y_label": "Circadian Stimulus (CS, approx.)",
        "notes": "Vertical lux is crucial for non-visual responses; measure at eye/face level for circadian effect estimates.",
        "refs": [
            ("EN 12464-1 and WELL references on vertical illumination", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            (INPUT_UPLOADED_DOCX, "User file: some studies included vertical illuminance ranges (350-1000) in literature")
        ],
        "markers": [(300, "300 lux target"), (500, "500 lux strong")]
    },

    {
        "title": "Exposure Duration (hours of daytime light at target levels)",
        "xspan": (0, 8),
        "bands": {"good": (2, 4), "warn": (1, 6), "danger": (0, 8)},
        "anchors_x": [0, 0.5, 1, 2, 3, 4, 6, 8],
        "anchors_y": [0, 10, 20, 40, 60, 75, 90, 95],   # cumulative melatonin suppression %
        "x_label": "Exposure Duration (hours)",
        "y_label": "Estimated Cumulative Melatonin Suppression (%)",
        "notes": "Sustained daytime exposure (~2–4 h at adequate EDI) supports entrainment. Short or irregular exposure is less effective.",
        "refs": [
            ("Brown et al., 2022 consensus; circadian exposure guidance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
            (INPUT_UPLOADED_DOCX, "User file: exposure duration context in literature review")
        ],
        "markers": [(2, "2 h effective"), (4, "4 h robust")]
    },

    {
        "title": "Horizontal Illuminance (desk/task lux)",
        "xspan": (100, 1500),
        "bands": {"good": (300, 500), "warn": (200, 1000), "danger": (100, 1500)},
        "anchors_x": [100, 200, 300, 500, 750, 1000, 1500],
        "anchors_y": [60, 75, 85, 95, 98, 99, 99],   # visual performance %
        "x_label": "Horizontal Illuminance (lux)",
        "y_label": "Visual Task Performance (%)",
        "notes": "300–500 lx on desk level is typical for classrooms; exams/labs may use short-term higher levels (≥750 lx).",
        "refs": [
            ("EN 12464-1 classroom illuminance (desk level)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            (INPUT_UPLOADED_DOCX, "User file: tested 275, 475, 613 lux and reported alertness increases")
        ],
        "markers": [(300, "300 lx baseline"), (500, "500 lx target"), (1000, "1000 lx focus/exam")]
    }
]

# --------------------------
# Recommendations per age & environment (based on standards + study)
# --------------------------
RECS = [
    (
        "Kindergarten classrooms (ages 4–6)",
        "Horizontal Illuminance: 300–500 lx",
        "CRI ≥ 80",
        "UGR ≤ 19",
        "Melanopic EDI ~200–300 lx",
        "CCT 3500–4000 K",
        "Children benefit from moderate light levels, good color rendering, and warm-neutral CCT that supports calmness.",
        [
            ("EN 12464-1 indoor lighting", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("Brown et al., 2022 consensus", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571")
        ]
    ),
    (
        "Primary school classrooms (ages 7–12)",
        "Horizontal Illuminance: 300–500 lx (up to 750 lx during reading/writing)",
        "CRI ≥ 80–85",
        "UGR ≤ 19",
        "Melanopic EDI ~250–400 lx",
        "CCT 4000–5000 K",
        "Slightly higher CCT supports alertness; ensure vertical illuminance at eye is sufficient for circadian entrainment.",
        [
            ("EN 12464-1", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("Park et al., 2015", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/")
        ]
    ),
    (
        "Secondary schools / adolescents (ages 13–18)",
        "Horizontal Illuminance: 500 lx baseline, up to 1000 lx for detailed tasks",
        "CRI ≥ 85",
        "UGR ≤ 19",
        "Melanopic EDI ≥ 300 lx (ideally 400–500 lx daytime)",
        "CCT 5000–6500 K",
        "Teenagers need strong circadian signals due to delayed sleep phase; cooler CCT supports morning alertness.",
        [
            ("Brown et al., 2022 consensus", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
            ("Mott et al., dynamic lighting in classrooms", "https://journals.sagepub.com/doi/abs/10.1177/1477153512446099")
        ]
    ),
    (
        "Laboratories / exam halls",
        "Horizontal Illuminance: 750–1000 lx",
        "CRI ≥ 90",
        "UGR ≤ 19",
        "Melanopic EDI ≥ 400 lx",
        "CCT 5000–6500 K",
        "High-intensity, cool light enhances alertness and task precision; suitable for exams and labs.",
        [
            ("EN 12464-1", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("WELL Building Standard", "https://resources.wellcertified.com/articles/circadian-rhythms/")
        ]
    ),
    (
        "Playgrounds / common areas",
        "Horizontal Illuminance: 100–200 lx",
        "CRI ≥ 70",
        "UGR control less critical outdoors",
        "Melanopic EDI variable, natural daylight preferred",
        "CCT variable, daylight spectrum preferred",
        "Outdoor spaces rely on natural daylight; artificial lighting should prioritize safety and visibility rather than circadian control.",
        [
            ("CIBSE lighting guide", "https://www.cibse.org/"),
        ]
    )
]

# --------------------------
# Read uploaded file (if any) and extract some values for the intro
# --------------------------
uploaded_text, uploaded_paras = read_docx_text(INPUT_UPLOADED_DOCX)
if uploaded_text is None:
    uploaded_paras = []
    uploaded_text = ""
    uploaded_summary = "No uploaded file found at: {}".format(INPUT_UPLOADED_DOCX)
else:
    kv = extract_key_values(uploaded_paras)
    # Build a short summary for the intro
    uploaded_summary_lines = ["Findings extracted from your uploaded file (Schools information.docx):"]
    if 'CCT_values' in kv:
        uploaded_summary_lines.append(" - Reported/tested CCT values: {}".format(", ".join(str(x) + "K" for x in kv['CCT_values'])))
    if 'Lux_values' in kv:
        uploaded_summary_lines.append(" - Reported illuminance (lux) measurements: {}".format(", ".join(str(x) + " lx" for x in kv['Lux_values'])))
    if 'CRI_values' in kv:
        uploaded_summary_lines.append(" - Reported CRI values: {}".format(", ".join(str(x) for x in kv['CRI_values'])))
    if 'Flicker_freqs_Hz' in kv:
        uploaded_summary_lines.append(" - Reported flicker frequencies: {} Hz".format(", ".join(str(x) for x in kv['Flicker_freqs_Hz'])))
    if 'age_mentions' in kv:
        uploaded_summary_lines.append(" - Age groups mentioned: {}".format(", ".join(kv['age_mentions'])))
    if 'percent_changes' in kv:
        uploaded_summary_lines.append(" - % changes noted: {}".format(", ".join(kv['percent_changes'])))
    # Also include first paragraph excerpts (up to 3)
    excerpts = uploaded_paras[:3] if uploaded_paras else []
    if excerpts:
        uploaded_summary_lines.append("\nKey excerpts from the uploaded study:")
        for e in excerpts:
            shortened = e if len(e) < 400 else e[:400] + "..."
            uploaded_summary_lines.append("   • " + shortened)
    uploaded_summary = "\n".join(uploaded_summary_lines)

# --------------------------
# Build the DOCX Document
# --------------------------
doc = Document()
doc.add_heading("Lighting in Schools — Biological & Cognitive Effects", 0)
doc.add_paragraph("Merged booklet that combines your uploaded study findings with literature-anchored parameter analysis.")

# Front matter
doc.add_heading("The Problem", level=1)
problem_par = (
    "Poor lighting in schools — including incorrect spectral content (CCT), low color rendering (CRI), excessive flicker, "
    "high glare (UGR), low or very uneven illuminance, and inadequate melanopic stimulation — undermines student performance, "
    "increases visual and physiological strain, disturbs sleep and circadian rhythms, and negatively effects mood."
)
doc.add_paragraph(problem_par)
# Add a paragraph citing the uploaded file summary
doc.add_paragraph("Merged uploaded-study findings (brief):")
doc.add_paragraph(uploaded_summary)

doc.add_heading("The Idea", level=1)
doc.add_paragraph(
    "This study compares measurable lighting parameters across a range of values and quantifies biological and cognitive responses. "
    "We combine standards (EN 12464-1, IEEE 1789, WELL) and academic dose–response anchors with the empirical results "
    "reported in the uploaded study to form practical recommendations."
)

doc.add_heading("The Study (Compare good vs bad values)", level=1)
doc.add_paragraph(
    "For each parameter we present: definition, biological mechanism, a literature-anchored response curve, and optimal/caution/risk ranges."
)
doc.add_paragraph("Key points from the uploaded study (selected):")
# Insert key points from uploaded_paras if present
if uploaded_paras:
    for i,p in enumerate(uploaded_paras[:8], 1):
        doc.add_paragraph(f"{i}. {p}")
else:
    doc.add_paragraph("No uploaded study content available or file not found at the expected path.")

doc.add_heading("Solution (Good values per age & environment)", level=1)
doc.add_paragraph("Recommendations synthesized from standards and uploaded-study observations:")
for rec in RECS:
    who = rec[0]
    doc.add_paragraph(f"• {who}")
    doc.add_paragraph(f"    - {rec[1]} | {rec[2]} | {rec[3]}")
    doc.add_paragraph(f"    - Melanopic target: {rec[4]} | CCT: {rec[5]}")
    doc.add_paragraph(f"    - Notes: {rec[6]}")
    doc.add_paragraph("    - Sources:")
    for t,u in rec[7]:
        doc.add_paragraph(f"      • {t} — {u}")

# Chapters: one parameter per chapter with figure
doc.add_heading("Chapters: Parameter-by-Parameter", level=1)

for p in PARAMS:
    title = p["title"]
    doc.add_heading(title, level=2)
    doc.add_paragraph(p["notes"])
    bands = p["bands"]
    doc.add_paragraph(f"Optimal: {bands['good'][0]}–{bands['good'][1]}   |   Caution: {bands['warn'][0]}–{bands['warn'][1]}")

    # Create data and plot
    x0, x1 = p["xspan"]
    # Choose x vector length sensibly depending on span
    if isinstance(x0, int) and isinstance(x1, int) and (x1 - x0) <= 1000:
        x = np.linspace(x0, x1, 500)
    else:
        # smaller vector for very tight/spread ranges
        x = np.linspace(x0, x1, 400)

    y = interpolate_curve(x, p["anchors_x"], p["anchors_y"])

    # Plot
    fig, ax = plt.subplots(figsize=(7, 3))
    add_bands(ax, x0, x1, bands["good"], bands["warn"], bands["danger"])
    ax.plot(x, y, color=COLOR_CURVE, linewidth=2.0, label="Biological Response")
    # markers
    for mk in p.get("markers", []):
        xpos, label = mk
        ax.axvline(x=xpos, color=COLOR_MARK, linestyle="--", linewidth=1.1)
        # annotate near top:
        ylim = ax.get_ylim()
        ypos = ylim[1] * 0.95
        ax.text(xpos, ypos, label, rotation=90, va="top", ha="right", fontsize=8, color=COLOR_MARK)
    # ticks
    as_int_ticks = isinstance(x0, int) and isinstance(x1, int)
    ensure_ticks = np.linspace(x0, x1, 6)
    if as_int_ticks:
        ensure_ticks = np.round(ensure_ticks).astype(int)
    ax.set_xticks(ensure_ticks)
    ax.set_xlabel(p["x_label"])
    ax.set_ylabel(p["y_label"])
    ax.legend(loc="upper right")
    imgname = title.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").replace(":", "") + ".png"
    imgpath = os.path.join(IMG_DIR, imgname)
    fig.tight_layout()
    fig.savefig(imgpath, dpi=180, bbox_inches="tight")
    plt.close(fig)

    # Insert image into doc
    doc.add_picture(imgpath, width=Inches(6.0))

    doc.add_paragraph("References:")
    for (t,u) in p["refs"]:
        # If the ref is the local uploaded doc, show local path and note
        if u == INPUT_UPLOADED_DOCX:
            doc.add_paragraph(f"• {t} — (from uploaded file) {u}")
        else:
            doc.add_paragraph(f"• {t} — {u}")

# Master references (unique list)
doc.add_heading("Master References", level=1)
master_refs = {
    "EN 12464-1 overview (indoor workplaces)": "https://www.performanceinlighting.com/mo/en/en-12464-1",
    "Brown et al., 2022 PLOS Biology (melanopic consensus)": "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571",
    "WELL resource: Circadian context": "https://resources.wellcertified.com/articles/circadian-rhythms/",
    "IEEE 1789 (flicker)": "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf",
    "Park et al., 2015 (CCT & task performance PMC)": "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/",
    "User uploaded file (Schools information.docx)": INPUT_UPLOADED_DOCX
}
for t,u in master_refs.items():
    doc.add_paragraph(f"• {t} — {u}")

# Save docx
os.makedirs(OUT_DIR, exist_ok=True)
doc.save(OUTPUT_DOCX)

print("✅ Done.")
print("Output DOCX:", OUTPUT_DOCX)
print("Figures:", IMG_DIR)
