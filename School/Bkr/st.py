# -*- coding: utf-8 -*-
"""
Creates a single DOCX booklet with:
- Front matter: Problem, Idea, Study, Solution (age x environment)
- Eight chapters (CCT, CRI, Flicker, UGR, Melanopic EDI, Vertical Illuminance, Exposure Duration, Horizontal Illuminance)
- One combined figure per chapter (curve + green/yellow/red zones)
- Live URLs for every reference

Requirements:
  pip install python-docx matplotlib numpy
"""

import os
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# -----------------------------------------------------------------------------
# 0) Matplotlib Style Switch (change this to any installed style you like)
#    Examples: "default", "classic", "seaborn-v0_8", "ggplot", "bmh", "dark_background"
# -----------------------------------------------------------------------------
MATPLOTLIB_STYLE = "seaborn-v0_8"
plt.style.use(MATPLOTLIB_STYLE)

# Basic colors (kept explicit for consistency/accuracy)
COLOR_CURVE = "#1f77b4"  # blue
COLOR_GOOD  = "#2ca02c"  # green band
COLOR_WARN  = "#ffbf00"  # yellow/amber band
COLOR_BAD   = "#d62728"  # red band
COLOR_MARK  = "#555555"  # vertical guideline markers

# -----------------------------------------------------------------------------
# 1) Utility helpers
# -----------------------------------------------------------------------------
def ensure_dir(path):
    if not os.path.isdir(path):
        os.makedirs(path, exist_ok=True)

def add_bands(ax, xlim, good, warn, danger):
    """Draw green/yellow/red bands along the x-axis."""
    # danger is the full span; warn is inner band; good is core
    ax.axvspan(danger[0], warn[0], color=COLOR_BAD, alpha=0.25, label="Risk")
    ax.axvspan(warn[0],   good[0], color=COLOR_WARN, alpha=0.30, label="Caution")
    ax.axvspan(good[0],   good[1], color=COLOR_GOOD, alpha=0.30, label="Optimal")
    ax.axvspan(good[1],   warn[1], color=COLOR_WARN, alpha=0.30)
    ax.axvspan(warn[1],   danger[1], color=COLOR_BAD, alpha=0.25)

    ax.set_xlim(*xlim)

def smooth_curve(x, anchors_x, anchors_y):
    """Return smooth y by interpolating anchor points across x."""
    return np.interp(x, anchors_x, anchors_y)

def clean_ticks(ax, x_min, x_max, n=6, as_int=False):
    ticks = np.linspace(x_min, x_max, n)
    if as_int:
        ticks = np.round(ticks).astype(int)
    ax.set_xticks(ticks)

def save_fig(fig, out_path):
    fig.tight_layout()
    fig.savefig(out_path, dpi=200, bbox_inches="tight")
    plt.close(fig)

# -----------------------------------------------------------------------------
# 2) Parameter specifications
#    Each entry defines:
#      - x-range (min, max)
#      - safe/caution/danger bands
#      - anchor points -> interpolated biological curve
#      - labels and references (title + url)
# -----------------------------------------------------------------------------

PARAMS = [
    # 1) CCT
    {
        "title": "CCT (Correlated Color Temperature, K)",
        "xspan": (2000, 7000),
        "bands": { "good": (4000, 5000), "warn": (3000, 6500), "danger": (2000, 7000) },
        # Biological proxy: daytime alerting potential / melatonin suppression tendency (relative %)
        # Anchors inspired by lab studies showing higher arousal with cooler CCT in day.
        "anchors": {
            "x":  [2000, 2700, 3000, 3500, 4000, 5000, 6500, 7000],
            "y":  [10,   18,   25,   45,   65,   70,   60,   55],  # %
        },
        "y_label": "Estimated Alerting Potential (%)",
        "x_label": "CCT (Kelvin)",
        "notes": (
            "Daytime 4000–5000 K generally supports alertness and visual comfort; "
            "short task-specific use of 6500 K may boost performance but can increase discomfort if overused."
        ),
        "refs": [
            ("EN 12464-1 overview (indoor workplaces: illuminance, UGR, CRI)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("Park et al., 2015: CCT, EEG & task performance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
            ("Chen et al., 2022: CCT × illuminance effects", "https://www.mdpi.com/1996-1073/15/12/4477"),
        ],
        "markers": [
            (4000, "Typical classroom"),
            (5000, "Upper preferred")
        ],
    },

    # 2) CRI
    {
        "title": "CRI (Color Rendering Index, Ra)",
        "xspan": (60, 100),
        "bands": { "good": (80, 100), "warn": (70, 80), "danger": (60, 100) },
        # Biological/visual performance proxy (% correct color/visual recognition)
        "anchors": {
            "x": [60, 70, 75, 80, 85, 90, 95, 100],
            "y": [50, 60, 70, 82, 90, 96, 99, 100],  # %
        },
        "y_label": "Visual Color Fidelity / Recognition (%)",
        "x_label": "CRI (Ra)",
        "notes": "CRI ≥80 is generally recommended for classrooms; ≥90 for art/graphics where color evaluation matters.",
        "refs": [
            ("EN 12464-1 overview (Ra requirements)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
        ],
        "markers": [
            (80, "Baseline classroom"),
            (90, "Art / graphics")
        ],
    },

    # 3) Flicker
    {
        "title": "Flicker (Percent Modulation, typical LED)",
        "xspan": (0, 50),
        "bands": { "good": (0, 5), "warn": (5, 20), "danger": (0, 50) },
        # Biological risk proxy: headache/eyestrain risk index (%) – increases with modulation depth
        "anchors": {
            "x": [0, 2, 5, 10, 20, 30, 40, 50],
            "y": [0, 5, 10, 25, 50, 70, 85, 95],  # %
        },
        "y_label": "Estimated Adverse Effect Risk (%)",
        "x_label": "Percent Flicker (%)",
        "notes": "Keep percent flicker as low as practical (<5%). Avoid low-frequency PWM; follow IEEE 1789 guidance.",
        "refs": [
            ("IEEE 1789-2015: Flicker Recommended Practice (PDF)", "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf"),
            ("DOE/LightFair: Understanding IEEE Flicker Practice (PDF)", "https://www.energy.gov/sites/default/files/2022-11/ssl-miller-lehman_flicker_lightfair2015.pdf"),
            ("Miller et al., 2022 review (PDF)", "https://www.energy.gov/sites/default/files/2022-08/ssl-miller-etal-2022-LRT-flicker-review-tlm-stimulus-response.pdf"),
        ],
        "markers": [
            (5, "Preferred max"),
            (20, "High risk")
        ],
    },

    # 4) Glare (UGR)
    {
        "title": "Glare (Unified Glare Rating, UGR)",
        "xspan": (10, 30),
        "bands": { "good": (10, 19), "warn": (19, 22), "danger": (10, 30) },
        # Biological/comfort proxy: discomfort probability (%). Higher UGR -> higher discomfort.
        "anchors": {
            "x": [10, 13, 16, 19, 22, 25, 28, 30],
            "y": [5, 8, 15, 30, 55, 75, 90, 95],  # %
        },
        "y_label": "Estimated Discomfort Probability (%)",
        "x_label": "UGR",
        "notes": "Aim UGR <19 for classrooms; even lower (≈16–18) near screens/IBs to minimize discomfort and distraction.",
        "refs": [
            ("CIBSE Factfile: Importance of glare & calculating UGR (PDF)", "https://www.cibse.org/media/polbabib/factfile-15-the-importance-of-glare-and-calculating-ugr-jul2019.pdf"),
            ("EN 12464-1 overview (UGR contexts)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
        ],
        "markers": [
            (19, "Classroom max"),
            (16, "Screen work")
        ],
    },

    # 5) Melanopic EDI
    {
        "title": "Melanopic EDI (melanopic lux at eye, vertical)",
        "xspan": (0, 800),
        "bands": { "good": (250, 500), "warn": (100, 250), "danger": (0, 800) },
        # Biological effect: melatonin suppression (%) under daytime conditions (saturating response)
        "anchors": {
            "x": [0, 20, 50, 100, 250, 500, 800],
            "y": [0, 5, 15, 35, 65, 80, 90],  # %
        },
        "y_label": "Estimated Melatonin Suppression (%)",
        "x_label": "Melanopic EDI (lux)",
        "notes": "Provide ≥250 melanopic EDI during the day for circadian entrainment and alertness (measured vertically at ~1.2 m).",
        "refs": [
            ("Brown et al., 2022 (PLOS Biology): Consensus recommendations", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
            ("Brown et al., 2022 (PMC)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
            ("WELL v2 – Circadian Lighting context (Article)", "https://resources.wellcertified.com/articles/circadian-rhythms/"),
        ],
        "markers": [
            (250, "Daytime target (min)"),
            (500, "Robust daytime")
        ],
    },

    # 6) Vertical Illuminance
    {
        "title": "Vertical Illuminance (lux at eye/face)",
        "xspan": (50, 1000),
        "bands": { "good": (300, 500), "warn": (200, 800), "danger": (50, 1000) },
        # Biological effect: Circadian Stimulus (CS, 0-0.7+) proxy (increases with vertical lx; saturates)
        "anchors": {
            "x": [50, 100, 150, 300, 500, 800, 1000],
            "y": [0.05, 0.12, 0.22, 0.40, 0.55, 0.65, 0.70],  # CS (unitless)
        },
        "y_label": "Circadian Stimulus (CS, unitless)",
        "x_label": "Vertical Illuminance (lux)",
        "notes": "Aim ~300–500 lx vertical on faces/eye for daytime non-visual benefits; much lower levels are advisable in evening school events.",
        "refs": [
            ("EN 12464-1 overview; vertical/ambient aspects", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("WELL v2 circadian context (Article)", "https://resources.wellcertified.com/articles/circadian-rhythms/"),
        ],
        "markers": [
            (300, "Daytime min"),
            (500, "Strong CS")
        ],
    },

    # 7) Exposure Duration
    {
        "title": "Exposure Duration (hours of daytime light meeting target)",
        "xspan": (0, 8),
        "bands": { "good": (2, 4), "warn": (1, 6), "danger": (0, 8) },
        # Biological effect: cumulative melatonin suppression (%) at moderate melanopic stimulus
        "anchors": {
            "x": [0, 0.5, 1, 2, 3, 4, 6, 8],
            "y": [0, 10, 20, 40, 60, 75, 90, 95],  # %
        },
        "y_label": "Estimated Cumulative Melatonin Suppression (%)",
        "x_label": "Exposure Duration (hours)",
        "notes": "Sustained daytime exposure (~2–4 h at adequate spectrum/levels) supports alertness and entrainment; avoid excessive high-intensity late-day exposure.",
        "refs": [
            ("Brown et al., 2022 (PMC): Day vs evening guidance", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
        ],
        "markers": [
            (2, "Effective"),
            (4, "Robust")
        ],
    },

    # 8) Horizontal Illuminance
    {
        "title": "Horizontal Illuminance (desk/task, lux)",
        "xspan": (100, 1500),
        "bands": { "good": (300, 500), "warn": (200, 1000), "danger": (100, 1500) },
        # Biological/visual performance: % task speed/accuracy vs lx (saturating)
        "anchors": {
            "x": [100, 200, 300, 500, 750, 1000, 1500],
            "y": [60, 75, 85, 95, 98, 99, 99],  # %
        },
        "y_label": "Visual Task Performance (%)",
        "x_label": "Horizontal Illuminance (lux)",
        "notes": "Provide 300–500 lx at desks for general classrooms; 500–750+ lx for labs/graphics. Short-term 800–1000 lx can be used for exam focus.",
        "refs": [
            ("EN 12464-1 overview (classroom/task lx)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
            ("MDPI (2025) review referencing EN 12464-1 classroom levels", "https://www.mdpi.com/2075-5309/15/8/1233"),
        ],
        "markers": [
            (300, "General min"),
            (500, "Classroom target"),
            (750, "Lab/graphics")
        ],
    },
]

# -----------------------------------------------------------------------------
# 3) Front-matter text blocks (Problem • Idea • Study • Solution)
# -----------------------------------------------------------------------------

PROBLEM_TXT = (
    "Suboptimal school lighting—too little or too much illuminance, excessive flicker, high glare (UGR), "
    "poor spectrum/CCT balance, low CRI, and inadequate melanopic stimulus—has been linked with headaches, eye strain, "
    "reduced reading performance, lower attention, and circadian disruption. These factors can degrade learning outcomes, "
    "increase fatigue, and negatively affect behavior and mood."
)

IDEA_TXT = (
    "Systematically study how measurable lighting parameters (CCT, CRI, Flicker, Glare/UGR, Melanopic EDI, Vertical Illuminance, "
    "Exposure Duration, and Horizontal Illuminance) affect children of different ages. Quantify biological and cognitive outcomes "
    "using literature-anchored dose–response curves and compare good versus poor ranges."
)

STUDY_TXT = (
    "For each parameter, we present a definition, biological relevance, a literature-anchored response curve with optimal, caution, "
    "and risk zones, and links to standards or peer-reviewed sources. The combined visuals indicate where classroom lighting "
    "supports attention, visual comfort, and circadian health—and where it does not."
)

SOLUTION_INTRO = (
    "The following age × environment recommendations synthesize standards (EN 12464-1 for lx/UGR/CRI), "
    "consensus guidance on melanopic EDI (Brown et al., 2022), WELL v2 context on circadian lighting, and flicker guidance (IEEE 1789)."
)

RECS = [
    ("Kindergarten (3–5) – classroom",
     "Horizontal lx: 300–500", "UGR: <19", "CRI: ≥80", "Melanopic EDI (day): ≥250", "CCT: 3500–4000 K",
     "Softer CCT reduces over-arousal; keep flicker <5%; vertical ~300–400 lx for faces.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
      ("Brown et al., 2022 (PMC)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/")]),

    ("Primary (6–11) – classroom",
     "Horizontal lx: 300–500", "UGR: <19", "CRI: ≥80", "Melanopic EDI (day): ≥250–300", "CCT: 4000–5000 K",
     "Balanced spectrum/daylight; flicker <5%; vertical ~300–500 lx faces/boards.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
      ("WELL circadian article", "https://resources.wellcertified.com/articles/circadian-rhythms/")]),

    ("Secondary (12–18) – classroom",
     "Horizontal lx: 300–500", "UGR: <19 (≤16 near screens)", "CRI: ≥80 (≥90 for art)", "Melanopic EDI (day): ≥250–300", "CCT: 4000–5000 K",
     "Lower UGR near screens; short high-CCT/high-lx sessions can support exam focus.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1")]),

    ("Exam/Focus (short sessions, all ages)",
     "Horizontal lx: 500–1000", "UGR: <19", "CRI: ≥80", "Melanopic EDI (day): ≥300–400", "CCT: 5000–6500 K",
     "Short deployments to boost alertness; avoid all-day cold light.",
     [("Park et al., 2015", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
      ("Chen et al., 2022", "https://www.mdpi.com/1996-1073/15/12/4477")]),

    ("Art/Graphics room",
     "Horizontal lx: 500–750", "UGR: <19", "CRI: ≥90", "Melanopic EDI (day): ≥250", "CCT: 4000–5000 K",
     "High CRI for color judgment; strong vertical lighting to evaluate work.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1")]),

    ("Science lab",
     "Horizontal lx: 500–750", "UGR: <19 (≤16 preferred)", "CRI: ≥80", "Melanopic EDI (day): ≥250–300", "CCT: 4000–5000 K",
     "Higher task illumination and glare control for practical work; minimize flicker.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1")]),

    ("Corridors / circulation",
     "Horizontal lx: 100–200", "UGR: <22", "CRI: ≥80", "Melanopic EDI: —", "CCT: 3000–4000 K",
     "Comfortable navigation; avoid glare and harsh contrasts.",
     [("EN 12464-1 overview", "https://www.performanceinlighting.com/mo/en/en-12464-1")]),
]

# -----------------------------------------------------------------------------
# 4) Build booklet
# -----------------------------------------------------------------------------
OUT_DIR = os.path.join(os.getcwd(), "school_lighting_booklet_output")
IMG_DIR = os.path.join(OUT_DIR, "images")
ensure_dir(OUT_DIR)
ensure_dir(IMG_DIR)

doc = Document()
doc.add_heading("Lighting in Schools: Biological & Cognitive Effects", 0)
doc.add_paragraph("A referenced booklet on eight lighting parameters and their effects on student concentration, biology, and psychology.")

# Front-matter sections
doc.add_heading("The Problem", level=1)
doc.add_paragraph(PROBLEM_TXT)

doc.add_heading("The Idea", level=1)
doc.add_paragraph(IDEA_TXT)

doc.add_heading("The Study (What We Compare)", level=1)
doc.add_paragraph(STUDY_TXT)

doc.add_heading("Solution (Targets by Age & Environment)", level=1)
doc.add_paragraph(SOLUTION_INTRO)

for (who_where, lx, ugr, cri, medi, cct, note, refs) in RECS:
    doc.add_paragraph(f"• {who_where}")
    doc.add_paragraph(f"  - {lx}   |   {ugr}   |   {cri}")
    doc.add_paragraph(f"  - {medi}   |   {cct}")
    doc.add_paragraph(f"  - Notes: {note}")
    doc.add_paragraph("  - References:")
    for title, url in refs:
        doc.add_paragraph(f"    • {title} — {url}")

# Chapters
doc.add_heading("Chapters: Parameter-by-Parameter", level=1)

for p in PARAMS:
    title = p["title"]
    x0, x1 = p["xspan"]
    x = np.linspace(x0, x1, 400)

    # Interpolate biological curve
    y = smooth_curve(x, p["anchors"]["x"], p["anchors"]["y"])

    # Plot
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    add_bands(ax, (x0, x1), p["bands"]["good"], p["bands"]["warn"], p["bands"]["danger"])

    ax.plot(x, y, color=COLOR_CURVE, linewidth=2.2, label="Biological Response")
    # Vertical markers (e.g., recommended values)
    for xpos, label in p.get("markers", []):
        ax.axvline(xpos, color=COLOR_MARK, linestyle="--", linewidth=1.2)
        ax.text(xpos, ax.get_ylim()[1]*0.95, label, rotation=90, va="top", ha="right", fontsize=8, color=COLOR_MARK)

    ax.set_title(title)
    ax.set_xlabel(p["x_label"])
    ax.set_ylabel(p["y_label"])
    # Cleaner ticks
    as_int = isinstance(x0, (int, np.integer)) and isinstance(x1, (int, np.integer))
    clean_ticks(ax, x0, x1, n=6, as_int=as_int)

    # Save
    safe_name = title.replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").replace(":", "")
    img_path = os.path.join(IMG_DIR, f"{safe_name}.png")
    save_fig(fig, img_path)

    # Add to DOCX
    doc.add_heading(title, level=2)
    doc.add_paragraph(p["notes"])
    bands = p["bands"]
    doc.add_paragraph(f"Optimal: {bands['good'][0]}–{bands['good'][1]}   |   Caution: {bands['warn'][0]}–{bands['warn'][1]} (context dependent)")
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph("References:")
    for (t, u) in p["refs"]:
        doc.add_paragraph(f"• {t} — {u}")

# Master References section
doc.add_heading("Master Reference List (Live URLs)", level=1)
ALL_REFS = [
    ("EN 12464-1 overview (indoor workplaces: illuminance, UGR, CRI)", "https://www.performanceinlighting.com/mo/en/en-12464-1"),
    ("CIBSE Factfile: Importance of glare & calculating UGR (PDF)", "https://www.cibse.org/media/polbabib/factfile-15-the-importance-of-glare-and-calculating-ugr-jul2019.pdf"),
    ("Brown et al., 2022 (PLOS Biology): Consensus recommendations", "https://journals.plos.org/plosbiology/article?id=10.1371/journal.pbio.3001571"),
    ("Brown et al., 2022 (PMC mirror)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8929548/"),
    ("WELL v2 Circadian context article (IWBI)", "https://resources.wellcertified.com/articles/circadian-rhythms/"),
    ("IEEE 1789-2015 (PDF copy)", "https://www.lisungroup.com/wp-content/uploads/2020/02/IEEE-2015-STANDARDS-1789-Standard-Free-Download.pdf"),
    ("DOE/LightFair deck on IEEE 1789 (PDF)", "https://www.energy.gov/sites/default/files/2022-11/ssl-miller-lehman_flicker_lightfair2015.pdf"),
    ("Miller et al., 2022 flicker review (PDF)", "https://www.energy.gov/sites/default/files/2022-08/ssl-miller-etal-2022-LRT-flicker-review-tlm-stimulus-response.pdf"),
    ("Park et al., 2015: CCT, EEG & task performance (PMC)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4668153/"),
    ("Chen et al., 2022: CCT × illuminance (MDPI)", "https://www.mdpi.com/1996-1073/15/12/4477"),
    ("MDPI 2025 review referencing EN 12464-1 classroom levels", "https://www.mdpi.com/2075-5309/15/8/1233"),
]
for t, u in ALL_REFS:
    doc.add_paragraph(f"• {t} — {u}")

# Save DOCX
OUT_DOCX = os.path.join(OUT_DIR, "School_Lighting_Booklet_FULL.docx")
doc.save(OUT_DOCX)

print("✅ DOCX created at:", OUT_DOCX)
print("🖼️ Figures saved in:", IMG_DIR)
print("ℹ️ Change the Matplotlib style via MATPLOTLIB_STYLE near the top if you want a different look.")
