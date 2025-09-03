# make_school_lighting_chapters.py
# Creates 9 separate DOCX files (one per parameter) with full detail:
# Definition, Ranges, Biological Effects (Hormones, Skin, Nervous system),
# Biochemical Pathways, Recommendations, Checklist, and Parameter-Specific References.

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTDIR = Path("./school_lighting_chapters")
OUTDIR.mkdir(parents=True, exist_ok=True)

# ---------- Style helpers ----------
def set_styles(doc):
    styles = doc.styles
    if "TitleLarge" not in styles:
        s = styles.add_style("TitleLarge", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Title"]
        s.font.name = "Calibri"; s.font.size = Pt(28); s.font.bold = True
    if "H1" not in styles:
        s = styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Heading 1"]
        s.font.name = "Calibri"; s.font.size = Pt(18); s.font.bold = True
    if "H2" not in styles:
        s = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Heading 2"]
        s.font.name = "Calibri"; s.font.size = Pt(14); s.font.bold = True
    if "Body" not in styles:
        s = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.font.name = "Calibri"; s.font.size = Pt(11)
    if "Bullet" not in styles:
        s = styles.add_style("Bullet", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["List Bullet"]
        s.font.name = "Calibri"; s.font.size = Pt(11)

def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph(title, style="TitleLarge")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        q = doc.add_paragraph(subtitle, style="Body")
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(doc, text): doc.add_paragraph(text, style="H1")
def add_h2(doc, text): doc.add_paragraph(text, style="H2")

def p(doc, text):
    for line in (text or "").split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip(), style="Body")

def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="Bullet")

def table2(doc, left, right):
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0,0).text, t.cell(0,1).text = "Optimal", left
    t.cell(1,0).text, t.cell(1,1).text = "Caution", right
    for c,w in enumerate([2.0, 4.5]):
        for r in range(2):
            t.cell(r,c).width = Inches(w)

def write_chapter(filename, chapter):
    doc = Document()
    set_styles(doc)
    add_title(doc, chapter["title"], chapter.get("subtitle"))
    add_h1(doc, "Definition");              p(doc, chapter["definition"])
    add_h1(doc, "Recommended Ranges");      table2(doc, chapter["ranges"]["optimal"], chapter["ranges"]["caution"])
    add_h1(doc, "Biological Effects")

    add_h2(doc, "Hormones (Endocrine)")
    bullets(doc, chapter["bio"]["hormones"])

    add_h2(doc, "Skin (Photobiology & Peripheral Clocks)")
    bullets(doc, chapter["bio"]["skin"])

    add_h2(doc, "Nervous System (ipRGC → SCN → CNS)")
    bullets(doc, chapter["bio"]["nervous"])

    add_h1(doc, "Biochemical Pathways (Mechanistic Detail)")
    bullets(doc, chapter["biochem"])

    add_h1(doc, "Classroom Recommendations")
    bullets(doc, chapter["recs"])

    add_h1(doc, "Quick Checklist")
    bullets(doc, chapter["checklist"])

    add_h1(doc, "References")
    bullets(doc, chapter["refs"])

    out = OUTDIR / filename
    doc.save(out)
    print(f"Saved: {out.resolve()}")

# ---------- Content for each chapter ----------
chapters = [
# 1) CCT
{
 "filename": "01_CCT.docx",
 "title": "CCT — Correlated Color Temperature (K)",
 "subtitle": "Spectral balance and non-visual biology in schools",
 "definition": "Descriptor of spectral appearance vs. a blackbody radiator. Higher CCT = blue-rich (shorter wavelengths); lower CCT = warm (longer wavelengths).",
 "ranges": {
   "optimal": "4000–5000 K (general instruction), 5000–6500 K (morning alertness/exams), 3000–3500 K (late-day calming).",
   "caution": "≤2700 K (daytime sleepiness risk) or ≥6500 K (discomfort/glare if uncontrolled)."
 },
 "bio": {
   "hormones": [
     "Blue-rich (~460–490 nm) → ipRGC (OPN4) activation → SCN → ↓ sympathetic tone to pineal → ↓ NE → ↓ AANAT → ↓ melatonin synthesis (day).",
     "Morning blue-enriched light supports CRH→ACTH→cortisol diurnal peak; stabilizes HPA rhythm.",
     "Daytime light increases serotonin turnover (raphe), supporting mood/attention; evening warm light permits melatonin rise."
   ],
   "skin": [
     "Typical classroom LEDs lack UVB → negligible vitamin D synthesis.",
     "Skin opsins (OPN3/OPN5) can entrain local clocks via G-protein–cAMP–CREB; systemic impact modest at indoor illuminances."
   ],
   "nervous": [
     "ipRGC glutamatergic input (NMDA) to SCN shifts CLOCK/BMAL1 → PER/CRY molecular clock phase.",
     "Blue light increases retinal dopamine, aiding contrast sensitivity and attentional performance."
   ],
 },
 "biochem": [
   "ipRGC→SCN: glutamate + PACAP → NMDA-dependent Ca²⁺ influx → CREB phosphorylation → Per1/Per2 transcription → phase shifts.",
   "SCN→PVN→IML→SCG→pineal: ↓ β-adrenergic signaling → ↓ cAMP/PKA → ↓ AANAT → melatonin↓.",
   "Serotonin (TPH2) day turnover; night: 5-HT → melatonin via AANAT/ASMT."
 ],
 "recs": [
   "Provide scene presets: Focus (5000–6500 K morning), General (4000–5000 K), Calm (3000–3500 K late-day).",
   "Coordinate CCT with glare control and vertical EDI targets."
 ],
 "checklist": [
   "CCT scenes mapped to schedule.",
   "Glare controlled when using higher CCT.",
   "Teacher control available."
 ],
 "refs": [
   "Park et al. (2015) — CCT & illuminance on performance.",
   "Brown TM et al. (2022) — Reporting light for non-visual effects; melanopic metrics.",
   "Mott et al. — Classroom dynamic lighting and reading fluency."
 ]
},

# 2) CRI
{
 "filename": "02_CRI.docx",
 "title": "CRI — Color Rendering Index (Ra)",
 "subtitle": "Color fidelity, strain, and indirect stress biology",
 "definition": "Fidelity of color appearance vs. a reference. High CRI improves accurate perception of materials/skin tones.",
 "ranges": {"optimal": "Ra ≥80 (classrooms), Ra ≥90 (art/labs).", "caution": "Ra 70–79 (non-critical areas only)."},
 "bio": {
   "hormones": [
     "Indirect effect: poor fidelity → visual discomfort/strain → sympathetic & HPA activation → cortisol↑ in susceptible students."
   ],
   "skin": [
     "No direct biochemical change; CRI is a fidelity metric, not dose of wavelengths."
   ],
   "nervous": [
     "Spectral gaps that degrade color constancy increase cortical load (V1/V4), promoting fatigue and reduced attention."
   ]
 },
 "biochem": [
   "Visual strain → sympathetic output (NE/Epi) → co-activates HPA (CRH→ACTH→cortisol).",
   "Retinal glutamate demand ↑ under difficult perception → ATP use and oxidative stress risk ↑."
 ],
 "recs": [
   "Specify Ra ≥80 for classrooms; ≥90 for labs/art.",
   "Avoid spectra with deep troughs affecting educational materials and skin tones."
 ],
 "checklist": [
   "CRI verified in luminaire data.",
   "Spot-check color charts at desk level."
 ],
 "refs": [
   "EN 12464-1 — Indoor workplaces (CRI guidance).",
   "Visual strain literature related to low-fidelity spectra."
 ]
},

# 3) Flicker
{
 "filename": "03_Flicker.docx",
 "title": "Flicker — Temporal Light Modulation",
 "subtitle": "Invisible flicker, comfort, and neural excitability",
 "definition": "Variation of light output over time; described by modulation %, frequency, and waveform. Can be imperceptible yet biologically active.",
 "ranges": {"optimal": "Percent modulation ≤5% across occupied dimming range; avoid low fundamentals.", "caution": "5–20% or fundamentals <100 Hz; evaluate stroboscopic risk."},
 "bio": {
   "hormones": [
     "Discomfort/stress from flicker → ↑ ACTH → cortisol↑; chronic exposure may destabilize HPA in sensitive individuals."
   ],
   "skin": ["No direct photochemical effect at classroom intensities."],
   "nervous": [
     "Low-frequency components can entrain abnormal cortical rhythms; trigger migraines/photosensitive seizures in vulnerable populations.",
     "Raises saccadic suppression demand → eye strain, headaches, reduced reading endurance."
   ]
 },
 "biochem": [
   "Repetitive drive → glutamatergic overactivation in visual cortex; excitotoxic susceptibility increases.",
   "Arousal circuits: locus coeruleus NE↑; HPA axis activation (CRH→ACTH→cortisol)."
 ],
 "recs": [
   "Specify drivers compliant with IEEE 1789; check flicker at multiple dim levels.",
   "Test under mains variation; avoid deep PWM at low frequencies."
 ],
 "checklist": [
   "Percent modulation and/or short-range index documented.",
   "No visible stroboscopic artifacts with moving objects."
 ],
 "refs": [
   "IEEE 1789 — Recommended practice for LED modulation (flicker).",
   "Clinical literature on photosensitive epilepsy/migraine triggers."
 ]
},

# 4) Glare
{
 "filename": "04_Glare_UGR.docx",
 "title": "Glare — Unified Glare Rating (UGR)",
 "subtitle": "Discomfort, visual fatigue, and stress pathways",
 "definition": "Discomfort arising from high luminance contrasts within the field of view, predicted by UGR (source luminance, position, background).",
 "ranges": {"optimal": "UGR ≤19 (classrooms).", "caution": "UGR 19–22 (caution), >22 (avoid)."},
 "bio": {
   "hormones": [
     "Persistent discomfort → sympathetic activation and HPA upregulation → cortisol↑."
   ],
   "skin": ["No direct skin effect."],
   "nervous": [
     "Retinal overstimulation → glutamate↑ → visual fatigue/headaches.",
     "Attention fragmentation from bright sources → working-memory efficiency↓."
   ]
 },
 "biochem": [
   "Aversive visual input engages limbic pathways (amygdala) → HPA activation.",
   "Photoreceptor bleaching/recovery cycles raise mitochondrial ROS; antioxidants (SOD, catalase) taxed."
 ],
 "recs": [
   "Use diffusers/microprismatic optics; avoid direct view of high-luminance emitters.",
   "Control reflected glare on boards/screens; coordinate luminance and CCT."
 ],
 "checklist": [
   "UGR verified in lighting calc.",
   "Check reflections from whiteboards and displays at student eye positions."
 ],
 "refs": [
   "EN 12464-1 / CIBSE LG — Glare limits.",
   "Studies linking glare to visual fatigue and task errors."
 ]
},

# 5) Horizontal Illuminance
{
 "filename": "05_Horizontal_Illuminance.docx",
 "title": "Horizontal Illuminance — Desk/Task (lx)",
 "subtitle": "Visual performance and non-visual support",
 "definition": "Illuminance on the working plane (desks). Adequate levels support reading speed, error reduction, and comfort.",
 "ranges": {"optimal": "300–500 lx general classrooms; 750–1000 lx short-term exams/labs (with glare control).", "caution": "200–299 lx (strain risk); >1000 lx (glare if uncontrolled)."},
 "bio": {
   "hormones": [
     "Higher daytime illuminance → stronger ipRGC drive → melatonin suppression; supports morning cortisol amplitude.",
     "Adequate light supports serotonin turnover and overall mood/attention."
   ],
   "skin": ["Indoor electric light (no UVB) → negligible vitamin D effect."],
   "nervous": [
     "Greater retinal drive → SCN stability → improved vigilance and executive function.",
     "Supports prefrontal dopamine tone, reducing errors and enhancing working memory."
   ]
 },
 "biochem": [
   "ipRGC glutamate/PACAP → NMDA-Ca²⁺→CREB→Per gene expression; SCN synchronizes peripheral clocks via AVP/VIP/GABA.",
   "Daylight components (when present) further reinforce circadian amplitude."
 ],
 "recs": [
   "Design for 300–500 lx at desks with uniformity ≥0.6.",
   "Use boost scenes (750–1000 lx) for exams; manage glare and flicker."
 ],
 "checklist": [
   "Lux measured across multiple desks and rows.",
   "Uniformity and contrast to board verified."
 ],
 "refs": [
   "EN 12464-1 — Classroom illuminance and uniformity.",
   "Park et al. — Illuminance and alertness/performance."
 ]
},

# 6) Vertical Illuminance
{
 "filename": "06_Vertical_Illuminance.docx",
 "title": "Vertical Illuminance — Eye-Level (lx)",
 "subtitle": "The better proxy for circadian stimulus",
 "definition": "Illuminance on a vertical plane at eye height; more predictive of non-visual responses than horizontal lux.",
 "ranges": {"optimal": "300–500 lx vertical (daytime).", "caution": "200–299 lx (weak circadian drive); >800 lx (check glare/UGR)."},
 "bio": {
   "hormones": [
     "Adequate vertical light efficiently suppresses melatonin by day; supports robust cortisol rhythm."
   ],
   "skin": ["Minimal direct effect absent UV; effects are retinally mediated."],
   "nervous": [
     "Triggers c-Fos in SCN; resets molecular clock (CLOCK/BMAL1→PER/CRY).",
     "Enhances locus coeruleus and basal forebrain activity → alertness, memory encoding."
   ]
 },
 "biochem": [
   "Light pulses → NMDA-dependent Ca²⁺ influx → CREB → Per1/Per2 expression → phase adjustment.",
   "SCN outputs modulate pineal AANAT via sympathetic chain."
 ],
 "recs": [
   "Measure vertical lx at student eye positions across the room.",
   "Combine with spectral tuning to meet melanopic targets (see mEDI)."
 ],
 "checklist": [
   "Vertical lx verified during morning hours.",
   "No direct view of high-luminance sources."
 ],
 "refs": [
   "WELL Building Standard — Vertical light at eye guidance.",
   "Brown TM et al. (2022) — Circadian-relevant measures."
 ]
},

# 7) Melanopic EDI
{
 "filename": "07_Melanopic_EDI.docx",
 "title": "Melanopic EDI — Equivalent Daylight Illuminance (melanopic lux)",
 "subtitle": "Spectrally weighted metric for ipRGC stimulus",
 "definition": "Photometric metric weighted to melanopsin sensitivity; better predictor of circadian/non-visual effects than photopic lux alone.",
 "ranges": {"optimal": "≥250–500 melanopic lx for students during daytime (especially morning).", "caution": "100–249 mEDI (weak); <100 mEDI (insufficient)."},
 "bio": {
   "hormones": [
     "Daytime ≥250 mEDI → robust melatonin suppression and entrainment; supports morning cortisol peak.",
     "Daytime light improves serotonin availability (precursor to nocturnal melatonin)."
   ],
   "skin": [
     "Skin opsins (e.g., OPN5) may align local circadian rhythms; systemic hormonal impact mostly retinally mediated."
   ],
   "nervous": [
     "Strong melanopic drive synchronizes SCN, improving arousal networks (noradrenergic/cholinergic)."
   ]
 },
 "biochem": [
   "OPN4 (Gq/11) → PLCβ → IP3/DAG → Ca²⁺ rise → transcriptional effects in SCN neurons.",
   "SCN coordinates peripheral oscillators via neuropeptides (VIP, AVP) stabilizing metabolism and cognition."
 ],
 "recs": [
   "Use spectrally tuned luminaires/daylight to reach morning mEDI targets.",
   "Verify with spectrometer or CIE S 026 calculator."
 ],
 "checklist": [
   "mEDI measured at eye height for seated students.",
   "Morning exposure window ≥2 h at target levels."
 ],
 "refs": [
   "Brown TM et al. (2022) — Reporting light for non-visual effects (melanopic metrics).",
   "CIE S 026/E:2018 — System for metrology of optical radiation for ipRGC-influenced responses.",
   "WELL — Circadian lighting feature guidance."
 ]
},

# 8) Uniformity
{
 "filename": "08_Uniformity.docx",
 "title": "Uniformity — Emin / Eavg",
 "subtitle": "Spatial consistency, comfort, and load on visual system",
 "definition": "Ratio of minimum to average illuminance. Higher uniformity means fewer dark corners and less adaptation stress.",
 "ranges": {"optimal": "≥0.6 in classrooms (≥0.7 desirable in exam halls if practicable).", "caution": "0.4–0.59 (caution); <0.4 (avoid)."},
 "bio": {
   "hormones": [
     "Uneven fields raise adaptation stress → sympathetic/HPA activation → cortisol↑ in susceptible students."
   ],
   "skin": ["Neutral at indoor levels."],
   "nervous": [
     "Frequent retinal adaptation (bleach/recover) increases metabolic load and visual cortex effort; attention stamina declines."
   ]
 },
 "biochem": [
   "Photoreceptor mitochondrial load↑ → ROS generation; antioxidant defenses (SOD, catalase) taxed.",
   "Chronic visual stress may upregulate inflammatory mediators (e.g., IL-6, TNF-α) in susceptible individuals."
 ],
 "recs": [
   "Lay out luminaires to minimize contrast; consider indirect components.",
   "Verify uniformity at desks and whiteboards with measurements."
 ],
 "checklist": [
   "Uniformity ratio from lighting calc documented.",
   "Spot measurements confirm design values."
 ],
 "refs": [
   "EN 12464-1 — Uniformity requirements for classrooms.",
   "Human factors studies on uneven lighting and visual stress."
 ]
},

# 9) Exposure Duration
{
 "filename": "09_Exposure_Duration.docx",
 "title": "Exposure Duration — Daily Light Dose",
 "subtitle": "Time × spectrum × intensity for robust entrainment",
 "definition": "Cumulative non-visual light exposure across the day. Both intensity and spectrum matter; morning/forenoon exposure is most impactful.",
 "ranges": {"optimal": "2–4 h/day of adequate vertical melanopic exposure (≥250 mEDI) in the morning/early afternoon.", "caution": "<2 h/day or irregular schedules (risk of weak entrainment/delayed sleep)."},
 "bio": {
   "hormones": [
     "Stable daily dose entrains melatonin onset and cortisol amplitude; supports mood and daytime alertness.",
     "Adequate daytime light supports serotonin synthesis → nighttime melatonin via AANAT/ASMT."
   ],
   "skin": [
     "If outdoor daylight is included: UVB converts 7-dehydrocholesterol → previtamin D3 → vitamin D3 (liver/kidney activation to calcitriol)."
   ],
   "nervous": [
     "SCN stabilization improves hippocampal LTP and memory consolidation; reduces daytime sleepiness/inattention."
   ]
 },
 "biochem": [
   "CLOCK/BMAL1 drive PER/CRY transcription; PER/CRY proteins inhibit their own activators (negative feedback). Light via SCN sets the phase.",
   "Vitamin D: skin cholecalciferol → 25(OH)D (liver) → 1,25(OH)₂D (kidney) → VDR-mediated gene transcription affecting immune/neural pathways."
 ],
 "recs": [
   "Schedule brightest/most blue-enriched scenes in first school hours; calmer/warmer scenes later.",
   "Encourage outdoor breaks when feasible to supplement daylight dose."
 ],
 "checklist": [
   "Morning light block achieved (≥2 h).",
   "Scene schedules mapped to timetable; holidays and seasonality considered."
 ],
 "refs": [
   "Brown TM et al. (2022) — Guidance on timing and reporting of non-visual light.",
   "Chronobiology literature: PER/CRY entrainment and cognitive outcomes."
 ]
},
]

# ---------- Generate all ----------
for ch in chapters:
    write_chapter(ch["filename"], ch)

print("\nAll chapters generated in:", OUTDIR.resolve())
