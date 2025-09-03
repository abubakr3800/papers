# make_intro_docx.py
# Creates a separate DOCX for the expanded Introduction

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTFILE = Path("./Intro_School_Lighting.docx")

def set_styles(doc):
    styles = doc.styles
    if "TitleLarge" not in styles:
        s = styles.add_style("TitleLarge", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Title"]
        s.font.name = "Calibri"; s.font.size = Pt(26); s.font.bold = True
    if "H1" not in styles:
        s = styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Heading 1"]
        s.font.name = "Calibri"; s.font.size = Pt(16); s.font.bold = True
    if "Body" not in styles:
        s = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.font.name = "Calibri"; s.font.size = Pt(11)

def add_title(doc, title):
    p = doc.add_paragraph(title, style="TitleLarge")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def p(doc, text): 
    doc.add_paragraph(text, style="Body")

# -------- Build the document --------
doc = Document()
set_styles(doc)

add_title(doc, "Introduction")

p(doc, 
"Lighting in schools has long been regarded primarily as a matter of visual comfort — "
"ensuring that students can read, write, and see the board without strain, while also meeting "
"energy efficiency requirements. Yet in recent decades, research in neuroscience, endocrinology, "
"and chronobiology has demonstrated that light is not only a visual input but also a biological signal. "
"The eye contains specialized photoreceptors (intrinsically photosensitive retinal ganglion cells, or ipRGCs) "
"that project to the brain’s master circadian clock in the suprachiasmatic nucleus (SCN). Through this pathway, "
"light regulates hormone secretion, sleep–wake timing, mood, and cognitive performance.")

doc.add_paragraph("The Problem", style="H1")
p(doc, "Traditional classroom lighting systems are optimized only for brightness and visibility, ignoring "
       "the non-visual biological effects of light. As a result, students are often exposed to lighting that is "
       "visually adequate but biologically disruptive. Key issues include:")
p(doc, "• Circadian disruption: high CCT or blue-rich light late in the day delays melatonin secretion.")
p(doc, "• Hormonal imbalance: insufficient vertical illuminance in the morning weakens cortisol amplitude.")
p(doc, "• Cognitive fatigue: poor uniformity, low CRI, and flicker induce strain and impaired attention.")
p(doc, "• Mood instability: inadequate melanopic stimulus reduces serotonin turnover.")
p(doc, "• Long-term risks: chronic disruption linked to metabolic, immune, and psychological disorders.")

doc.add_paragraph("The Idea", style="H1")
p(doc, "The central idea of this framework is that light can be described and controlled through measurable "
       "parameters — CCT, CRI, flicker, glare, horizontal and vertical illuminance, melanopic EDI, uniformity, "
       "and exposure duration. By aligning these parameters with their biological, hormonal, skin, nervous system, "
       "and biochemical effects, lighting can be designed not just for seeing but for learning and wellbeing.")

doc.add_paragraph("Side Effects of Poor Lighting", style="H1")
p(doc, "Ignoring biological effects produces consequences beyond discomfort, including disrupted circadian alignment, "
       "abnormal melatonin suppression, cortisol flattening, headaches, reduced serotonin synthesis, and lower classroom engagement.")

doc.add_paragraph("Our Solution", style="H1")
p(doc, "This booklet provides a parameter-based framework that integrates biology with classroom lighting design. "
       "For each parameter, we present definitions, recommended ranges, biological effects, biochemical pathways, "
       "recommendations, and checklists. By shifting from a purely visual model to a biological + visual model, "
       "schools can create environments that enhance attention, stabilize circadian rhythms, protect long-term health, "
       "and ultimately improve educational outcomes.")

# Save
doc.save(OUTFILE)
print(f"Saved introduction file: {OUTFILE.resolve()}")
