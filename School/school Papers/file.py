from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
import os

# Create a new document
doc = Document()

# Title
doc.add_heading("The Effect of Classroom Lighting on Students’ Concentration, Biology, and Psychology", level=0)

# Section: Problem
doc.add_heading("1. Problem", level=1)
doc.add_paragraph(
    "Poor classroom lighting has significant negative effects on children’s health, psychology, and learning outcomes. "
    "Inadequate or excessive light can cause visual discomfort, headaches, reduced concentration, sleep disturbances, "
    "and long-term effects on circadian rhythms. Studies show that incorrect lighting parameters such as low CRI, high flicker, "
    "and improper CCT can worsen attention spans and reduce academic performance."
)
doc.add_paragraph("Reference: https://doi.org/10.1016/j.buildenv.2019.106219")

# Section: Idea
doc.add_heading("2. Idea", level=1)
doc.add_paragraph(
    "The idea of this study is to simulate and analyze the impact of different lighting parameters on students of various ages, "
    "and determine the biological and psychological effects of good versus poor lighting conditions. "
    "By comparing recommended values to harmful values, we aim to propose lighting strategies that improve concentration, "
    "well-being, and overall academic performance."
)
doc.add_paragraph("Reference: https://www.mdpi.com/2076-3417/11/18/8661")

# Section: Study
doc.add_heading("3. Study (Parameters and Their Effects)", level=1)

# Define parameter ranges and effects
parameters = {
    "CCT (Color Temperature, K)": {"ranges": [(2700, 3000, "Too warm → sleepiness", "red"),
                                              (3500, 5000, "Optimal alertness", "green"),
                                              (6000, 8000, "Too cold → eye strain", "yellow")]},
    "CRI (Color Rendering Index)": {"ranges": [(0, 70, "Poor color rendering", "red"),
                                               (80, 100, "Good visual quality", "green")]} ,
    "Flicker (%)": {"ranges": [(20, 100, "Dangerous flicker", "red"),
                               (5, 20, "Noticeable flicker", "yellow"),
                               (0, 5, "Imperceptible", "green")]} ,
    "Glare (UGR)": {"ranges": [(22, 30, "High glare → discomfort", "red"),
                               (19, 22, "Tolerable", "yellow"),
                               (10, 19, "Comfortable", "green")]} ,
    "Melanopic EDI (lux)": {"ranges": [(0, 100, "Too low → circadian disruption", "red"),
                                       (250, 300, "Optimal", "green"),
                                       (500, 1000, "Over-stimulating", "yellow")]} ,
    "Vertical Illuminance (lux)": {"ranges": [(0, 200, "Too dim", "red"),
                                              (300, 500, "Good", "green"),
                                              (800, 1200, "Too bright", "yellow")]} ,
    "Exposure Duration (hours)": {"ranges": [(0, 2, "Insufficient exposure", "red"),
                                             (4, 8, "Balanced", "green"),
                                             (10, 14, "Over-exposure", "yellow")]} ,
    "Horizontal Illuminance (Lux)": {"ranges": [(0, 200, "Too dim", "red"),
                                                (300, 500, "Good", "green"),
                                                (800, 1500, "Too bright → glare", "yellow")]} ,
}

# Create graphs for each parameter
graph_files = []
for param, details in parameters.items():
    plt.figure(figsize=(6,2))
    for start, end, label, color in details["ranges"]:
        plt.axvspan(start, end, color=color, alpha=0.3, label=label)
    plt.title(param)
    plt.xlabel("Value")
    plt.yticks([])
    handles, labels = plt.gca().get_legend_handles_labels()
    by_label = dict(zip(labels, handles))
    plt.legend(by_label.values(), by_label.keys(), loc="upper right")
    file_name = f"/mnt/data/{param.replace(' ', '_').replace('(', '').replace(')', '').replace('/', '_')}.png"
    plt.savefig(file_name, bbox_inches="tight")
    plt.close()
    graph_files.append((param, file_name))

    # Insert into DOCX
    doc.add_heading(param, level=2)
    doc.add_picture(file_name, width=Inches(5))
    doc.add_paragraph(f"Effect of {param} on students.")

# Section: Solution
doc.add_heading("4. Solution (Recommended Values)", level=1)
doc.add_paragraph(
    "Based on research findings, the following values are recommended for school environments:\n"
    "- CCT: 4000–5000 K in classrooms\n"
    "- CRI: ≥80\n"
    "- Flicker: <5%\n"
    "- Glare (UGR): <19\n"
    "- Melanopic EDI: ~250–300 lux\n"
    "- Vertical Illuminance: 300–500 lux\n"
    "- Exposure Duration: 4–8 hours of balanced daylight/artificial light\n"
    "- Horizontal Illuminance: 300–500 lux\n"
)
doc.add_paragraph("References: https://doi.org/10.1016/j.buildenv.2019.106219 , https://www.mdpi.com/2076-3417/11/18/8661")

# Save DOCX
output_path = "/mnt/data/School_Lighting_Booklet_Final.docx"
doc.save(output_path)

output_path
